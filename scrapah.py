import requests
from bs4 import BeautifulSoup
import pandas as pd
import time

def search_jobs(query, location, num_pages=1):
    jobs = []
    headers = {'User-Agent': 'Mozilla/5.0'}
    for page in range(num_pages):
        url = f"https://www.indeed.com/jobs?q={query}&l={location}&start={page*10}"
        response = requests.get(url, headers=headers)
        time.sleep(1)  # Be polite and avoid overwhelming the server
        soup = BeautifulSoup(response.text, 'html.parser')
        for card in soup.find_all('div', class_='job_seen_beacon'):
            title = card.find('h2', class_='jobTitle').text.strip()
            company = card.find('span', class_='companyName').text.strip()
            location = card.find('div', class_='companyLocation').text.strip()
            summary = card.find('div', class_='job-snippet').text.strip()
            link_tag = card.find('a', href=True)
            link = 'https://www.indeed.com' + link_tag['href'] if link_tag else ''
            jobs.append({
                'Title': title,
                'Company': company,
                'Location': location,
                'Summary': summary,
                'Link': link
            })
    return pd.DataFrame(jobs)

# Example usage
df = search_jobs('Full Stack Developer', 'Chicago, IL', num_pages=2)
print(df.head())


import spacy

nlp = spacy.load('en_core_web_sm')

def extract_keywords(text):
    doc = nlp(text)
    keywords = [token.lemma_ for token in doc if token.pos_ in ['NOUN', 'VERB', 'ADJ'] and not token.is_stop]
    return list(set(keywords))

# Example usage
job_description = df.iloc[0]['Summary']
job_keywords = extract_keywords(job_description)
print(job_keywords)


from docx import Document

def update_resume(resume_path, job_keywords):
    doc = Document(resume_path)
    # Simple example: Add keywords to the end of the resume
    doc.add_heading('Relevant Skills', level=1)
    doc.add_paragraph(', '.join(job_keywords))
    tailored_resume_path = 'Tailored_Resume.docx'
    doc.save(tailored_resume_path)
    return tailored_resume_path

# Example usage
tailored_resume = update_resume('Your_Resume.docx', job_keywords)


import datetime
from docx import Document

def generate_cover_letter(template_path, applicant_info, job_info):
    with open(template_path, 'r') as file:
        template = file.read()

    # Combine skills into a string
    skills = ', '.join(job_info['skills'])

    # Prepare the content by replacing placeholders
    cover_letter_content = template.format(
        your_name=applicant_info['name'],
        your_address=applicant_info['address'],
        your_city_state_zip=applicant_info['city_state_zip'],
        your_email=applicant_info['email'],
        your_phone=applicant_info['phone'],
        date=datetime.datetime.now().strftime('%B %d, %Y'),
        hiring_manager_name=job_info.get('hiring_manager_name', 'Hiring Manager'),
        company_name=job_info['company'],
        company_address=job_info.get('company_address', ''),
        company_city_state_zip=job_info.get('company_city_state_zip', ''),
        job_title=job_info['job_title'],
        skills=skills
    )

    # Save the cover letter as a Word document
    tailored_cover_letter_path = f"Cover_Letter_{job_info['company']}_{job_info['job_title']}.docx"
    doc = Document()
    for line in cover_letter_content.strip().split('\n'):
        doc.add_paragraph(line)
    doc.save(tailored_cover_letter_path)
    return tailored_cover_letter_path

def prepare_application_data(df, resume_path, cover_letter_template, applicant_info):
    applications = []
    for index, row in df.iterrows():
        job_keywords = extract_keywords(row['Summary'])
        tailored_resume = update_resume(resume_path, job_keywords)
        job_info = prepare_job_info(row, job_keywords)
        tailored_cover_letter = generate_cover_letter(
            cover_letter_template,
            applicant_info,
            job_info
        )
        applications.append({
            'Company': row['Company'],
            'Job Title': row['Title'],
            'Link': row['Link'],
            'Tailored Resume': tailored_resume,
            'Tailored Cover Letter': tailored_cover_letter
        })
    applications_df = pd.DataFrame(applications)
    applications_df.to_csv('Applications.csv', index=False)
    print("Application data prepared and saved to Applications.csv.")


applicant_info = {
    'name': 'Robert Neyrinck',
    'address': '1209 N State Pkwy',
    'city_state_zip': 'Chicago, IL 60610',
    'email': 'robert.a.neyrinck@gmail.com',
    'phone': '(872) 333-7804',
}
def prepare_job_info(row, job_keywords):
    job_info = {
        'company': row['Company'],
        'job_title': row['Title'],
        'skills': job_keywords,
        # Optionally add 'hiring_manager_name', 'company_address', 'company_city_state_zip' if available
    }
    return job_info

# Example usage
# Assuming df is your DataFrame of job postings
prepare_application_data(df, 'Your_Resume.docx', 'Cover_Letter_Template.txt', applicant_info)



