import requests
import pandas as pd
import spacy
from docx import Document
import datetime
import time
import random
from bs4 import BeautifulSoup

# Load spaCy language model
nlp = spacy.load('en_core_web_sm')

def get_public_ip():
    try:
        response = requests.get('https://api.ipify.org')
        if response.status_code == 200:
            return response.text
        else:
            print("Unable to get public IP address.")
            return None
    except Exception as e:
        print(f"Error getting public IP: {e}")
        return None

def search_jobs_careerjet(query, location, page=1, pagesize=10):
    api_url = 'http://public.api.careerjet.net/search'
    user_ip = get_public_ip()
    user_agent = 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)'  # Example user agent

    if not user_ip:
        print("Public IP address is required to make the API request.")
        return pd.DataFrame()

    params = {
        'locale_code': 'en_US',
        'keywords': query,
        'location': location,
        'page': page,
        'pagesize': pagesize,
        'affid': '50f6b5fcca95ad3283da9025659d0ae2',  # Replace with your actual affiliate ID
        'user_ip': user_ip,
        'user_agent': user_agent,
    }
    try:
        response = requests.get(api_url, params=params, timeout=10)
        response.raise_for_status()
        data = response.json()
        if 'type' in data and data['type'] == 'ERROR':
            print(f"API Error: {data.get('error', 'Unknown error')}")
            return pd.DataFrame()
        jobs = []
        for job in data.get('jobs', []):
            jobs.append({
                'Title': job.get('title'),
                'Company': job.get('company'),
                'Location': job.get('locations'),
                'Summary': job.get('description'),
                'Link': job.get('url'),
            })
        return pd.DataFrame(jobs)
    except requests.exceptions.RequestException as e:
        print(f"An error occurred: {e}")
        return pd.DataFrame()

def search_jobs_careerjet_multiple_pages(query, location, start_page=1, end_page=1, pagesize=10):
    all_jobs = []
    for page in range(start_page, end_page + 1):
        df = search_jobs_careerjet(query, location, page=page, pagesize=pagesize)
        all_jobs.append(df)
        time.sleep(random.uniform(1, 3))  # Wait between 1 to 3 seconds
    if all_jobs:
        combined_df = pd.concat(all_jobs, ignore_index=True)
        combined_df.drop_duplicates(subset=['Link'], inplace=True)
        return combined_df
    else:
        return pd.DataFrame()

def clean_text(text):
    # Remove HTML tags
    soup = BeautifulSoup(text, 'html.parser')
    cleaned_text = soup.get_text(separator=' ')
    # Remove extra whitespace
    cleaned_text = ' '.join(cleaned_text.split())
    return cleaned_text

def extract_keywords(text):
    # Predefined set of relevant skills
    skill_set = set([
        'Python', 'JavaScript', 'TypeScript', 'Node.js', 'React', 'Django', 'AWS', 'GCP', 'Azure',
        'Docker', 'Kubernetes', 'Terraform', 'CI/CD', 'DevOps', 'Microservices', 'PostgreSQL', 'Redis',
        'AI', 'Machine Learning', 'REST', 'RESTful', 'APIs', 'Agile', 'Scrum', 'Kanban', 'Git', 'Linux', 'NoSQL',
        'Automation', 'Cloud', 'Infrastructure', 'IaC', 'Full Stack', 'Backend', 'Frontend', 'Leadership',
        'Mentoring', 'Team Lead', 'Software Development', 'System Architecture', 'CI/CD Pipelines', 'SQL',
        'Jenkins', 'Ansible', 'Puppet', 'Chef', 'Flask', 'FastAPI', 'Express.js', 'MongoDB', 'GraphQL'
    ])
    # Tokenize and normalize text
    doc = nlp(text)
    text_tokens = set(token.text for token in doc)
    # Find intersection of skills and text tokens
    keywords = list(skill_set.intersection(text_tokens))
    return keywords

def prepare_job_info(row, job_keywords):
    job_info = {
        'company': row['Company'],
        'job_title': row['Title'],
        'skills': job_keywords,
        # Optionally add 'hiring_manager_name', 'company_address', 'company_city_state_zip' if available
    }
    return job_info

def update_resume(resume_path, job_keywords, job_info):
    doc = Document(resume_path)
    # Find 'Skills' section and update it
    skills_found = False
    for i, paragraph in enumerate(doc.paragraphs):
        if 'Skills' in paragraph.text:
            skills_found = True
            # Update the next paragraph with new skills
            if i + 1 < len(doc.paragraphs):
                existing_skills = doc.paragraphs[i + 1].text
                # Combine existing skills with new ones
                all_skills = set(existing_skills.split(', ')) | set(job_keywords)
                doc.paragraphs[i + 1].text = ', '.join(all_skills)
            else:
                # Add new paragraph with skills
                doc.add_paragraph(', '.join(job_keywords))
            break
    if not skills_found:
        # Add a 'Skills' section if it doesn't exist
        doc.add_heading('Skills', level=1)
        doc.add_paragraph(', '.join(job_keywords))
    # Save with a unique filename
    safe_company = job_info['company'].replace(' ', '_').replace('/', '_')
    safe_title = job_info['job_title'].replace(' ', '_').replace('/', '_')
    tailored_resume_path = f"Tailored_Resume_{safe_company}_{safe_title}.docx"
    doc.save(tailored_resume_path)
    return tailored_resume_path

def generate_cover_letter(template_path, applicant_info, job_info):
    with open(template_path, 'r') as file:
        template = file.read()

    # Prepare a natural sentence with skills
    if job_info['skills']:
        if len(job_info['skills']) == 1:
            skills_sentence = f"My expertise in {job_info['skills'][0]} aligns closely with the requirements for the {job_info['job_title']} role."
        else:
            skills_list = ', '.join(job_info['skills'][:-1])
            skills_sentence = f"My expertise in {skills_list}, and {job_info['skills'][-1]} aligns closely with the requirements for the {job_info['job_title']} role."
    else:
        skills_sentence = f"My expertise aligns closely with the requirements for the {job_info['job_title']} role."

    # Prepare the content by replacing placeholders
    cover_letter_content = template.format(
        your_name=applicant_info['name'],
        your_address=applicant_info['address'],
        your_city_state_zip=applicant_info['city_state_zip'],
        your_email=applicant_info['email'],
        your_phone=applicant_info['phone'],
        date=datetime.datetime.now().strftime('%B %d, %Y'),
        hiring_manager_name=job_info.get('hiring_manager_name', 'Hiring Manager'),
        company_name=job_info['company'],
        company_address=job_info.get('company_address', ''),
        company_city_state_zip=job_info.get('company_city_state_zip', ''),
        job_title=job_info['job_title'],
        skills_sentence=skills_sentence
    )

    # Save the cover letter as a Word document
    safe_company = job_info['company'].replace(' ', '_').replace('/', '_')
    safe_title = job_info['job_title'].replace(' ', '_').replace('/', '_')
    tailored_cover_letter_path = f"Cover_Letter_{safe_company}_{safe_title}.docx"
    doc = Document()
    for line in cover_letter_content.strip().split('\n'):
        doc.add_paragraph(line)
    doc.save(tailored_cover_letter_path)
    return tailored_cover_letter_path

def prepare_application_data(df, resume_path, cover_letter_template, applicant_info):
    applications = []
    for index, row in df.iterrows():
        # Clean the job summary
        cleaned_summary = clean_text(row['Summary'])
        # Extract keywords from the job description
        job_keywords = extract_keywords(cleaned_summary)
        job_info = prepare_job_info(row, job_keywords)
        # Generate tailored resume and cover letter
        tailored_resume = update_resume(resume_path, job_keywords, job_info)
        tailored_cover_letter = generate_cover_letter(
            cover_letter_template,
            applicant_info,
            job_info
        )
        # Store application data, including the extracted skills and default status
        applications.append({
            'Company': row['Company'],
            'Job Title': row['Title'],
            'Link': row['Link'],
            'Tailored Resume': tailored_resume,
            'Tailored Cover Letter': tailored_cover_letter,
            'Skills': ', '.join(job_keywords),
            'Status': 'Not Applied',  # Add default status
        })
    # Save all applications to Applications.csv
    applications_df = pd.DataFrame(applications)
    applications_df.to_csv('Applications.csv', index=False)
    print("Application data prepared and saved to Applications.csv.")


# Applicant Information
applicant_info = {
    'name': 'Robert Neyrinck',
    'address': '1209 N State Pkwy',
    'city_state_zip': 'Chicago, IL 60610',
    'email': 'robert.a.neyrinck@gmail.com',
    'phone': '(872) 333-7804',
    'websites': [
        'https://curved-slash-06a.notion.site/Frontpage-Coming-Soon-10afc8851dcd809f81c9ce044dabd0b0',
        'https://rneyrinck.github.io/Portfolio-2022/'
    ],
    'education': 'Certificate SWE, General Assembly',
    'social_profiles': [
        'LinkedIn: https://linkedin.com/in/robert-neyrinck',
        'GitHub: https://github.com/rneyrinck'
    ],
}

# Main Execution
if __name__ == "__main__":
    # Retrieve jobs from multiple pages
    df = search_jobs_careerjet_multiple_pages(
        'Full Stack Developer',
        'Chicago, IL',
        start_page=1,
        end_page=2,  # Adjust as needed
        pagesize=10
    )

    print(f"Total jobs retrieved: {len(df)}")

    # Prepare application materials
    prepare_application_data(df, 'robert neyrinck resume.docx', 'Cover_Letter_Template.txt', applicant_info)
