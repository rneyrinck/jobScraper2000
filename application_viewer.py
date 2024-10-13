import sys
import os
import json
import pandas as pd
import webbrowser
import requests
import time
import random
import datetime
import spacy
import traceback
import urllib.parse
from bs4 import BeautifulSoup
from docx import Document
from PyQt5.QtWidgets import (
    QApplication, QWidget, QVBoxLayout, QHBoxLayout, QTableWidget, QTableWidgetItem,
    QPushButton, QLabel, QMessageBox, QTextEdit, QScrollArea, QComboBox, QMenuBar, QAction,
    QFileDialog, QLineEdit, QSpinBox, QDialogButtonBox, QDialog, QProgressDialog
)
from PyQt5.QtCore import Qt, QMimeData, QUrl
from PyQt5.QtGui import QDrag

# Load spaCy language model
try:
    nlp = spacy.load('en_core_web_sm')
except OSError:
    # If the model is not found, prompt to download it
    print("spaCy model 'en_core_web_sm' not found. Downloading...")
    from spacy.cli import download

    download('en_core_web_sm')
    nlp = spacy.load('en_core_web_sm')

# Define status options and their sort order
status_options = [
    'Not Applied',
    'Applied',
    'Interview Scheduled',
    'Interviewed',
    'Offer Received',
    'Accepted Offer',
    'Rejected',
    'No Response',
]

status_sort_order = {
    'Not Applied': 0,
    'Applied': 1,
    'Interview Scheduled': 2,
    'Interviewed': 3,
    'Offer Received': 4,
    'Accepted Offer': 5,
    'Rejected': 6,
    'No Response': 7,
}

# Define your base skills (same as skill_set in extract_keywords)
base_skills = [
    'Python', 'JavaScript', 'TypeScript', 'Node.js', 'React', 'Django', 'AWS', 'GCP', 'Azure',
    'Docker', 'Kubernetes', 'Terraform', 'CI/CD', 'DevOps', 'Microservices', 'PostgreSQL', 'Redis',
    'AI', 'Machine Learning', 'REST', 'RESTful', 'APIs', 'Agile', 'Scrum', 'Kanban', 'Git', 'Linux', 'NoSQL',
    'Automation', 'Cloud', 'Infrastructure', 'IaC', 'Full Stack', 'Backend', 'Frontend', 'Leadership',
    'Mentoring', 'Team Lead', 'Software Development', 'System Architecture', 'CI/CD Pipelines', 'SQL',
    'Jenkins', 'Ansible', 'Puppet', 'Chef', 'Flask', 'FastAPI', 'Express.js', 'MongoDB', 'GraphQL'
]

# Applicant Information
applicant_info = {
    'name': 'Robert Neyrinck',
    'address': '1209 N State Pkwy',
    'city_state_zip': 'Chicago, IL 60610',
    'email': 'robert.a.neyrinck@gmail.com',
    'phone': '(872) 333-7804',
    'websites': [
        'https://curved-slash-06a.notion.site/Frontpage-Coming-Soon-10afc8851dcd809f81c9ce044dabd0b0',
        'https://rneyrinck.github.io/Portfolio-2022/'
    ],
    'education': 'Certificate SWE, General Assembly',
    'social_profiles': [
        'LinkedIn: https://linkedin.com/in/robert-neyrinck',
        'GitHub: https://github.com/rneyrinck'
    ],
}


class ApplicationViewer(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle('Applications Viewer')
        self.resize(1000, 600)
        self.sort_column = 'Status'
        self.sort_order = True  # True for ascending order
        self.current_row = None
        self.current_application = None
        self.load_config()
        self.create_widgets()
        self.load_data()

    def load_config(self):
        try:
            with open('config.json', 'r') as f:
                config = json.load(f)
                self.resume_template_path = config.get('resume_template_path', '')
        except FileNotFoundError:
            self.resume_template_path = ''

    def save_config(self):
        config = {
            'resume_template_path': self.resume_template_path,
            # Add other configurations if needed
        }
        with open('config.json', 'w') as f:
            json.dump(config, f)

    def create_widgets(self):
        layout = QHBoxLayout()
        # Create Menu Bar
        self.menu_bar = QMenuBar(self)
        file_menu = self.menu_bar.addMenu('File')
        edit_menu = self.menu_bar.addMenu('Edit')

        # Add "Generate Applications" action
        generate_applications_action = QAction('Generate Applications', self)
        generate_applications_action.triggered.connect(self.generate_applications)
        file_menu.addAction(generate_applications_action)

        # Add "Upload Resume" action
        upload_resume_action = QAction('Upload Resume', self)
        upload_resume_action.triggered.connect(self.upload_resume)
        file_menu.addAction(upload_resume_action)

        # Add "Edit Resume Template" action
        edit_resume_action = QAction('Edit Resume Template', self)
        edit_resume_action.triggered.connect(self.edit_resume_template)
        edit_menu.addAction(edit_resume_action)

        # Main layout adjustments
        main_layout = QVBoxLayout()
        main_layout.setMenuBar(self.menu_bar)
        main_layout.addLayout(layout)
        self.setLayout(main_layout)

        # Create the table to display applications
        self.table = QTableWidget()
        self.table.setColumnCount(3)
        self.table.setHorizontalHeaderLabels(['Company', 'Job Title', 'Status'])
        self.table.cellClicked.connect(self.on_cell_clicked)
        self.table.setSelectionBehavior(QTableWidget.SelectRows)
        self.table.setSortingEnabled(False)  # We will handle sorting manually
        self.table.horizontalHeader().sectionClicked.connect(self.on_header_clicked)
        layout.addWidget(self.table)

        # Create a scroll area for the details and buttons
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        layout.addWidget(scroll_area)

        # Create a widget to hold the scroll area's content
        scroll_content = QWidget()
        scroll_area.setWidget(scroll_content)

        # Create a layout for the details and buttons
        details_layout = QVBoxLayout(scroll_content)

        # Labels to display details
        self.company_label = QLabel('Company: ')
        self.title_label = QLabel('Job Title: ')
        details_layout.addWidget(self.company_label)
        details_layout.addWidget(self.title_label)

        # Status Dropdown
        self.status_label = QLabel('Status:')
        details_layout.addWidget(self.status_label)
        self.status_dropdown = QComboBox()
        self.status_dropdown.addItems(status_options)
        self.status_dropdown.currentIndexChanged.connect(self.status_changed)
        self.status_dropdown.setEnabled(False)
        details_layout.addWidget(self.status_dropdown)

        # Button to open job link
        self.link_button = QPushButton('Open Job Link')
        self.link_button.clicked.connect(self.open_link)
        self.link_button.setEnabled(False)
        details_layout.addWidget(self.link_button)

        # Buttons for tailored resume and cover letter
        self.resume_button = QPushButton('Drag Resume')
        self.resume_button.pressed.connect(self.drag_resume)
        self.resume_button.setEnabled(False)
        details_layout.addWidget(self.resume_button)

        self.cover_button = QPushButton('Drag Cover Letter')
        self.cover_button.pressed.connect(self.drag_cover_letter)
        self.cover_button.setEnabled(False)
        details_layout.addWidget(self.cover_button)

        # Preview Buttons
        self.preview_resume_button = QPushButton('Preview Resume')
        self.preview_resume_button.clicked.connect(self.preview_resume)
        self.preview_resume_button.setEnabled(False)
        details_layout.addWidget(self.preview_resume_button)

        self.preview_cover_button = QPushButton('Preview Cover Letter')
        self.preview_cover_button.clicked.connect(self.preview_cover_letter)
        self.preview_cover_button.setEnabled(False)
        details_layout.addWidget(self.preview_cover_button)

        # Add a separator
        separator = QLabel('-' * 40)
        details_layout.addWidget(separator)

        # Skills Section
        self.skills_label = QLabel('Job-Specific Skills:')
        details_layout.addWidget(self.skills_label)
        self.skills_text = QTextEdit()
        self.skills_text.setReadOnly(True)
        self.skills_text.setFixedHeight(60)
        details_layout.addWidget(self.skills_text)
        self.copy_job_skills_button = QPushButton('Copy Job-Specific Skills')
        self.copy_job_skills_button.clicked.connect(self.copy_job_skills)
        self.copy_job_skills_button.setEnabled(False)
        details_layout.addWidget(self.copy_job_skills_button)

        # Base Skills Section
        self.base_skills_label = QLabel('Your Base Skills:')
        details_layout.addWidget(self.base_skills_label)
        self.base_skills_text = QTextEdit()
        self.base_skills_text.setReadOnly(True)
        self.base_skills_text.setFixedHeight(60)
        self.base_skills_text.setText(', '.join(base_skills))
        details_layout.addWidget(self.base_skills_text)
        self.copy_base_skills_button = QPushButton('Copy Base Skills')
        self.copy_base_skills_button.clicked.connect(self.copy_base_skills)
        details_layout.addWidget(self.copy_base_skills_button)

        # Combined Skills Section
        self.copy_combined_skills_button = QPushButton('Copy Combined Skills')
        self.copy_combined_skills_button.clicked.connect(self.copy_combined_skills)
        self.copy_combined_skills_button.setEnabled(False)
        details_layout.addWidget(self.copy_combined_skills_button)

        # Websites Section
        self.websites_label = QLabel('Websites:')
        details_layout.addWidget(self.websites_label)
        self.websites_layout = QVBoxLayout()
        details_layout.addLayout(self.websites_layout)
        self.create_website_widgets()

        # Education Section
        self.education_label = QLabel('Education:')
        details_layout.addWidget(self.education_label)
        self.education_text = QTextEdit()
        self.education_text.setReadOnly(True)
        self.education_text.setText(applicant_info.get('education', ''))
        self.education_text.setFixedHeight(40)
        details_layout.addWidget(self.education_text)
        self.copy_education_button = QPushButton('Copy Education to Clipboard')
        self.copy_education_button.clicked.connect(self.copy_education)
        details_layout.addWidget(self.copy_education_button)

        # Social Profiles Section
        self.social_label = QLabel('Social Profiles:')
        details_layout.addWidget(self.social_label)
        self.social_layout = QVBoxLayout()
        details_layout.addLayout(self.social_layout)
        self.create_social_widgets()

        # Add stretch to push widgets to the top
        details_layout.addStretch()

    def create_website_widgets(self):
        # Clear any existing widgets
        self.clear_layout(self.websites_layout)
        websites = applicant_info.get('websites', [])
        for website in websites:
            h_layout = QHBoxLayout()
            label = QLabel(f'<a href="{website}">{website}</a>')
            label.setOpenExternalLinks(True)
            copy_button = QPushButton('Copy')
            copy_button.clicked.connect(lambda checked, text=website: self.copy_to_clipboard(text))
            h_layout.addWidget(label)
            h_layout.addWidget(copy_button)
            self.websites_layout.addLayout(h_layout)

    def create_social_widgets(self):
        # Clear any existing widgets
        self.clear_layout(self.social_layout)
        social_profiles = applicant_info.get('social_profiles', [])
        for profile in social_profiles:
            h_layout = QHBoxLayout()
            if ': ' in profile:
                label_text, url = profile.split(': ', 1)
                label = QLabel(f'{label_text}: <a href="{url}">{url}</a>')
            else:
                label = QLabel(f'<a href="{profile}">{profile}</a>')
            label.setOpenExternalLinks(True)
            copy_button = QPushButton('Copy')
            copy_button.clicked.connect(lambda checked, text=profile: self.copy_to_clipboard(text))
            h_layout.addWidget(label)
            h_layout.addWidget(copy_button)
            self.social_layout.addLayout(h_layout)

    def clear_layout(self, layout):
        if layout is not None:
            while layout.count():
                child = layout.takeAt(0)
                if child.widget():
                    child.widget().deleteLater()
                elif child.layout():
                    self.clear_layout(child.layout())

    def load_data(self):
        try:
            self.df = pd.read_csv('Applications.csv')
            # Ensure 'Status' column exists
            if 'Status' not in self.df.columns:
                self.df['Status'] = 'Not Applied'
            self.df.reset_index(drop=True, inplace=True)  # Reset index to ensure alignment
            # Add 'StatusOrder' column
            self.df['StatusOrder'] = self.df['Status'].map(status_sort_order)
            # Sort by 'StatusOrder' by default
            self.df.sort_values(by=['StatusOrder'], inplace=True)
            self.df.reset_index(drop=True, inplace=True)  # Reset index after sorting
            # Populate the table
            self.populate_table()
        except (FileNotFoundError, pd.errors.EmptyDataError):
            self.df = pd.DataFrame(columns=['Company', 'Job Title', 'Status'])
            self.populate_table()

    def populate_table(self):
        self.table.setRowCount(len(self.df))
        for index, row in self.df.iterrows():
            self.table.setItem(index, 0, QTableWidgetItem(str(row.get('Company', ''))))
            self.table.setItem(index, 1, QTableWidgetItem(str(row.get('Job Title', ''))))
            self.table.setItem(index, 2, QTableWidgetItem(str(row.get('Status', 'Not Applied'))))

    def on_cell_clicked(self, row, column):
        try:
            self.current_row = row
            self.current_application = self.df.iloc[row]
            company = str(self.current_application.get('Company', ''))
            job_title = str(self.current_application.get('Job Title', ''))
            self.company_label.setText(f"Company: {company}")
            self.title_label.setText(f"Job Title: {job_title}")
            self.link_button.setEnabled(True)
            self.resume_button.setEnabled(True)
            self.cover_button.setEnabled(True)
            self.preview_resume_button.setEnabled(True)
            self.preview_cover_button.setEnabled(True)

            # Update skills
            skills = str(self.current_application.get('Skills', ''))
            self.skills_text.setText(skills)
            self.copy_job_skills_button.setEnabled(True)
            self.copy_combined_skills_button.setEnabled(True)

            # Set status in the dropdown without triggering the signal
            status = str(self.current_application.get('Status', 'Not Applied'))
            index = self.status_dropdown.findText(status)
            if index >= 0:
                self.status_dropdown.blockSignals(True)
                self.status_dropdown.setCurrentIndex(index)
                self.status_dropdown.blockSignals(False)
            else:
                self.status_dropdown.blockSignals(True)
                self.status_dropdown.setCurrentIndex(0)  # Default to 'Not Applied'
                self.status_dropdown.blockSignals(False)

            self.status_dropdown.setEnabled(True)
        except Exception as e:
            QMessageBox.critical(self, "Error", f"An error occurred: {e}")

    def on_header_clicked(self, logicalIndex):
        # Column names corresponding to table columns
        column_names = ['Company', 'Job Title', 'Status']
        if logicalIndex >= len(column_names):
            return
        sort_column = column_names[logicalIndex]
        # Toggle sort order (ascending/descending)
        if self.sort_column == sort_column:
            self.sort_order = not self.sort_order
        else:
            self.sort_column = sort_column
            self.sort_order = True  # Default to ascending
        # Perform sorting
        if sort_column == 'Status':
            # Ensure 'StatusOrder' column is up to date
            self.df['StatusOrder'] = self.df['Status'].map(status_sort_order)
            self.df.sort_values(by=['StatusOrder'], ascending=self.sort_order, inplace=True)
        else:
            self.df.sort_values(by=[sort_column], ascending=self.sort_order, inplace=True)
        self.df.reset_index(drop=True, inplace=True)
        # Repopulate the table
        self.populate_table()

    def status_changed(self, index):
        if hasattr(self, 'current_application'):
            new_status = self.status_dropdown.currentText()
            print(f"Changing status to: {new_status} for row: {self.current_row}")  # Debugging
            # Temporarily block signals to prevent recursive calls
            self.status_dropdown.blockSignals(True)
            # Update the DataFrame with the new status
            self.df.at[self.current_row, 'Status'] = new_status
            # Update the 'StatusOrder' for sorting
            self.df.at[self.current_row, 'StatusOrder'] = status_sort_order.get(new_status, 0)
            # Save the updated DataFrame to Applications.csv (excluding 'StatusOrder')
            columns_to_save = [col for col in self.df.columns if col != 'StatusOrder']
            self.df.to_csv('Applications.csv', index=False, columns=columns_to_save)
            # Update the specific cell in the table to reflect the new status
            self.table.setItem(self.current_row, 2, QTableWidgetItem(str(new_status)))
            # Re-enable signals
            self.status_dropdown.blockSignals(False)
            # Display a confirmation message to the user
            QMessageBox.information(self, "Status Updated", f"Application status updated to '{new_status}'.")

    def open_link(self):
        if hasattr(self, 'current_application'):
            url = self.current_application.get('Link', '')
            print(f"Opening URL: {url}")  # Debugging
            if url:
                webbrowser.open(url)
            else:
                QMessageBox.warning(self, "Warning", "No link available for this application.")

    def drag_resume(self):
        if hasattr(self, 'current_application'):
            file_path = self.current_application.get('Tailored Resume', '')
            self.start_drag(file_path)

    def drag_cover_letter(self):
        if hasattr(self, 'current_application'):
            file_path = self.current_application.get('Tailored Cover Letter', '')
            self.start_drag(file_path)

    def start_drag(self, file_path):
        if file_path and os.path.exists(file_path):
            mime_data = QMimeData()
            url = QUrl.fromLocalFile(os.path.abspath(file_path))
            mime_data.setUrls([url])

            drag = QDrag(self)
            drag.setMimeData(mime_data)
            drag.exec_(Qt.CopyAction)
        else:
            QMessageBox.warning(self, "Error", f"File not found: {file_path}")

    def copy_job_skills(self):
        skills = self.skills_text.toPlainText()
        clipboard = QApplication.clipboard()
        clipboard.setText(skills)
        QMessageBox.information(self, "Copied", "Job-specific skills copied to clipboard.")

    def copy_base_skills(self):
        skills = self.base_skills_text.toPlainText()
        clipboard = QApplication.clipboard()
        clipboard.setText(skills)
        QMessageBox.information(self, "Copied", "Base skills copied to clipboard.")

    def copy_combined_skills(self):
        job_skills = set(self.skills_text.toPlainText().split(', '))
        base_skills_set = set(self.base_skills_text.toPlainText().split(', '))
        combined_skills = ', '.join(sorted(job_skills.union(base_skills_set)))
        clipboard = QApplication.clipboard()
        clipboard.setText(combined_skills)
        QMessageBox.information(self, "Copied", "Combined skills copied to clipboard.")

    def copy_education(self):
        text = self.education_text.toPlainText()
        clipboard = QApplication.clipboard()
        clipboard.setText(text)
        QMessageBox.information(self, "Copied", "Education copied to clipboard.")

    def copy_to_clipboard(self, text):
        clipboard = QApplication.clipboard()
        clipboard.setText(text)
        QMessageBox.information(self, "Copied", f"'{text}' copied to clipboard.")

    def upload_resume(self):
        options = QFileDialog.Options()
        options |= QFileDialog.ReadOnly
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Select Resume File",
            "",
            "Word Documents (*.docx);;All Files (*)",
            options=options
        )
        if file_path:
            self.resume_template_path = file_path
            QMessageBox.information(self, "Resume Uploaded", f"Resume template updated to {file_path}")
            # Save the path to a config file
            self.save_config()

    def edit_resume_template(self):
        # Load the resume template
        if os.path.exists(self.resume_template_path):
            # Display the resume content in a text editor
            self.resume_editor = ResumeEditor(self.resume_template_path)
            self.resume_editor.show()
        else:
            QMessageBox.warning(self, "Error", f"Resume template not found at {self.resume_template_path}")

    def preview_resume(self):
        if hasattr(self, 'current_application'):
            file_path = self.current_application.get('Tailored Resume', '')
            if os.path.exists(file_path):
                self.preview_document(file_path, 'Resume Preview')
            else:
                QMessageBox.warning(self, "Error", f"File not found: {file_path}")

    def preview_cover_letter(self):
        if hasattr(self, 'current_application'):
            file_path = self.current_application.get('Tailored Cover Letter', '')
            if os.path.exists(file_path):
                self.preview_document(file_path, 'Cover Letter Preview')
            else:
                QMessageBox.warning(self, "Error", f"File not found: {file_path}")

    def preview_document(self, file_path, title):
        try:
            doc = Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            content = '\n'.join(full_text)
            # Display content in a new window
            self.doc_viewer = DocumentViewer(content, title)
            self.doc_viewer.show()
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to load document: {e}")

    def generate_applications(self):
        # Check if resume template is available
        if not os.path.exists(self.resume_template_path):
            reply = QMessageBox.question(
                self,
                "Resume Template Not Found",
                "Resume template not found. Would you like to upload one?",
                QMessageBox.Yes | QMessageBox.No
            )
            if reply == QMessageBox.Yes:
                self.upload_resume()
                if not os.path.exists(self.resume_template_path):
                    QMessageBox.warning(self, "Operation Cancelled",
                                        "Cannot generate applications without a resume template.")
                    return
            else:
                QMessageBox.warning(self, "Operation Cancelled",
                                    "Cannot generate applications without a resume template.")
                return

        # Collect job search parameters
        dialog = JobSearchDialog()
        if dialog.exec_() == QDialog.Accepted:
            # Ensure dialog is closed
            QApplication.processEvents()
            # Proceed to generate applications
            self.run_application_generation(dialog.keywords, dialog.location, dialog.start_page, dialog.end_page,
                                            dialog.page_size)
        else:
            QMessageBox.information(self, "Operation Cancelled", "Job search cancelled by the user.")

    def run_application_generation(self, keywords, location, start_page, end_page, page_size):
        # Create a progress dialog
        progress_dialog = QProgressDialog("Generating applications...", None, 0, 0, self)
        progress_dialog.setWindowTitle("Please Wait")
        progress_dialog.setWindowModality(Qt.WindowModal)
        progress_dialog.setCancelButton(None)
        progress_dialog.show()
        QApplication.processEvents()

        try:
            # Retrieve jobs from multiple pages
            df = search_jobs_careerjet_multiple_pages(
                keywords,
                location,
                start_page=start_page,
                end_page=end_page,
                pagesize=page_size,
                progress_dialog=progress_dialog
            )

            if df.empty:
                progress_dialog.close()
                QMessageBox.warning(self, "No Jobs Found", "No jobs were found with the given parameters.")
                return

            # Update progress dialog for application processing
            total_applications = len(df)
            progress_dialog.setMaximum(total_applications)
            progress_dialog.setLabelText("Preparing application materials...")
            QApplication.processEvents()

            # Prepare application materials
            prepare_application_data(df, self.resume_template_path, 'Cover_Letter_Template.txt', applicant_info,
                                     progress_dialog=progress_dialog)
            self.load_data()

            progress_dialog.close()
            QMessageBox.information(self, "Success", f"Generated applications for {len(df)} jobs.")
        except Exception as e:
            progress_dialog.close()
            traceback.print_exc()
            QMessageBox.critical(self, "Error", f"An error occurred: {e}")


class ResumeEditor(QWidget):
    def __init__(self, resume_path):
        super().__init__()
        self.setWindowTitle('Edit Resume Template')
        self.resume_path = resume_path
        self.create_widgets()
        self.load_resume()

    def create_widgets(self):
        layout = QVBoxLayout()
        self.setLayout(layout)
        self.text_edit = QTextEdit()
        layout.addWidget(self.text_edit)

        # Save Button
        self.save_button = QPushButton('Save')
        self.save_button.clicked.connect(self.save_resume)
        layout.addWidget(self.save_button)

    def load_resume(self):
        # Load the resume content
        try:
            doc = Document(self.resume_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            self.text_edit.setPlainText('\n'.join(full_text))
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to load resume: {e}")

    def save_resume(self):
        # Save the content back to the resume file
        try:
            content = self.text_edit.toPlainText()
            doc = Document()
            for line in content.split('\n'):
                doc.add_paragraph(line)
            doc.save(self.resume_path)
            QMessageBox.information(self, "Success", "Resume template saved successfully.")
            self.close()
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to save resume: {e}")


class DocumentViewer(QWidget):
    def __init__(self, content, title):
        super().__init__()
        self.setWindowTitle(title)
        self.create_widgets(content)

    def create_widgets(self, content):
        layout = QVBoxLayout()
        self.setLayout(layout)
        self.text_edit = QTextEdit()
        self.text_edit.setReadOnly(True)
        self.text_edit.setPlainText(content)
        layout.addWidget(self.text_edit)


class JobSearchDialog(QDialog):
    def __init__(self):
        super().__init__()
        self.setWindowTitle('Job Search Parameters')
        self.keywords = None
        self.location = None
        self.start_page = None
        self.end_page = None
        self.page_size = None
        self.create_widgets()

    def create_widgets(self):
        layout = QVBoxLayout()
        self.setLayout(layout)

        # Job Keywords
        self.keywords_label = QLabel('Job Keywords:')
        self.keywords_input = QLineEdit('Full Stack Developer')
        layout.addWidget(self.keywords_label)
        layout.addWidget(self.keywords_input)

        # Location
        self.location_label = QLabel('Location:')
        self.location_input = QLineEdit('Chicago, IL')
        layout.addWidget(self.location_label)
        layout.addWidget(self.location_input)

        # Start Page
        self.start_page_label = QLabel('Start Page:')
        self.start_page_input = QSpinBox()
        self.start_page_input.setMinimum(1)
        self.start_page_input.setValue(1)
        layout.addWidget(self.start_page_label)
        layout.addWidget(self.start_page_input)

        # End Page
        self.end_page_label = QLabel('End Page:')
        self.end_page_input = QSpinBox()
        self.end_page_input.setMinimum(1)
        self.end_page_input.setValue(2)
        layout.addWidget(self.end_page_label)
        layout.addWidget(self.end_page_input)

        # Page Size
        self.page_size_label = QLabel('Page Size:')
        self.page_size_input = QSpinBox()
        self.page_size_input.setMinimum(1)
        self.page_size_input.setValue(10)
        layout.addWidget(self.page_size_label)
        layout.addWidget(self.page_size_input)

        # Buttons
        self.button_box = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        self.button_box.accepted.connect(self.on_accept)
        self.button_box.rejected.connect(self.reject)
        layout.addWidget(self.button_box)

    def on_accept(self):
        self.keywords = self.keywords_input.text()
        self.location = self.location_input.text()
        self.start_page = self.start_page_input.value()
        self.end_page = self.end_page_input.value()
        self.page_size = self.page_size_input.value()
        self.accept()  # This will close the dialog and set the result code to Accepted

    def reject(self):
        self.keywords = None
        self.close()


# Data Preparation Functions

def get_public_ip():
    try:
        response = requests.get('https://api.ipify.org')
        if response.status_code == 200:
            return response.text
        else:
            print("Unable to get public IP address.")
            return None
    except Exception as e:
        print(f"Error getting public IP: {e}")
        return None


def search_jobs_careerjet(query, location, page=1, pagesize=10):
    api_url = 'http://public.api.careerjet.net/search'
    user_ip = get_public_ip()
    user_agent = 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)'  # Example user agent

    if not user_ip:
        print("Public IP address is required to make the API request.")
        return pd.DataFrame()

    params = {
        'locale_code': 'en_US',
        'keywords': query,
        'location': location,
        'page': page,
        'pagesize': pagesize,
        'affid': '50f6b5fcca95ad3283da9025659d0ae2',  # Your affiliate ID
        'user_ip': user_ip,
        'user_agent': user_agent,
    }
    try:
        response = requests.get(api_url, params=params, timeout=10)
        response.raise_for_status()
        data = response.json()
        if 'type' in data and data['type'] == 'ERROR':
            print(f"API Error: {data.get('error', 'Unknown error')}")
            return pd.DataFrame()
        jobs = []
        for job in data.get('jobs', []):
            original_url = job.get('url')
            encoded_url = urllib.parse.quote_plus(original_url)
            affiliate_link = f"https://www.careerjet.com/xyz/affid/{params['affid']}?dest_url={encoded_url}"
            print(f"Job URL: {affiliate_link}")  # Debugging line to check the URL
            jobs.append({
                'Title': job.get('title'),
                'Company': job.get('company'),
                'Location': job.get('locations'),
                'Summary': job.get('description'),
                'Link': affiliate_link,
            })
        return pd.DataFrame(jobs)
    except requests.exceptions.RequestException as e:
        print(f"An error occurred: {e}")
        return pd.DataFrame()


def search_jobs_careerjet_multiple_pages(query, location, start_page=1, end_page=1, pagesize=10, progress_dialog=None):
    all_jobs = []
    total_pages = end_page - start_page + 1
    current_progress = 0
    for page in range(start_page, end_page + 1):
        df = search_jobs_careerjet(query, location, page=page, pagesize=pagesize)
        all_jobs.append(df)
        time.sleep(random.uniform(1, 3))  # Wait between 1 to 3 seconds

        if progress_dialog:
            current_progress += 1
            progress_dialog.setValue(current_progress)
            progress_dialog.setLabelText(f"Fetching page {page}/{end_page}")
            QApplication.processEvents()

    if all_jobs:
        combined_df = pd.concat(all_jobs, ignore_index=True)
        combined_df.drop_duplicates(subset=['Link'], inplace=True)
        return combined_df
    else:
        return pd.DataFrame()


def clean_text(text):
    # Remove HTML tags
    soup = BeautifulSoup(text, 'html.parser')
    cleaned_text = soup.get_text(separator=' ')
    # Remove extra whitespace
    cleaned_text = ' '.join(cleaned_text.split())
    return cleaned_text


def extract_keywords(text):
    # Predefined set of relevant skills
    skill_set = set(base_skills)
    # Tokenize and normalize text
    doc = nlp(text)
    text_tokens = set(token.text for token in doc)
    # Find intersection of skills and text tokens
    keywords = list(skill_set.intersection(text_tokens))
    return keywords


def prepare_job_info(row, job_keywords):
    job_info = {
        'company': row['Company'],
        'job_title': row['Title'],
        'skills': job_keywords,
        # Optionally add 'hiring_manager_name', 'company_address', 'company_city_state_zip' if available
    }
    return job_info


def update_resume(resume_path, job_keywords, job_info):
    doc = Document(resume_path)
    # Find 'Skills' section and update it
    skills_found = False
    for i, paragraph in enumerate(doc.paragraphs):
        if 'Skills' in paragraph.text:
            skills_found = True
            # Update the next paragraph with new skills
            if i + 1 < len(doc.paragraphs):
                existing_skills = doc.paragraphs[i + 1].text
                # Combine existing skills with new ones
                all_skills = set(existing_skills.split(', ')) | set(job_keywords)
                doc.paragraphs[i + 1].text = ', '.join(all_skills)
            else:
                # Add new paragraph with skills
                doc.add_paragraph(', '.join(job_keywords))
            break
    if not skills_found:
        # Add a 'Skills' section if it doesn't exist
        doc.add_heading('Skills', level=1)
        doc.add_paragraph(', '.join(job_keywords))
    # Save with a unique filename in the tailored_documents directory
    safe_company = job_info['company'].replace(' ', '_').replace('/', '_')
    safe_title = job_info['job_title'].replace(' ', '_').replace('/', '_')
    os.makedirs('tailored_documents', exist_ok=True)
    tailored_resume_path = os.path.join('tailored_documents', f"Tailored_Resume_{safe_company}_{safe_title}.docx")
    doc.save(tailored_resume_path)
    return tailored_resume_path


def generate_cover_letter(template_path, applicant_info, job_info):
    with open(template_path, 'r') as file:
        template = file.read()

    # Prepare a natural sentence with skills
    if job_info['skills']:
        if len(job_info['skills']) == 1:
            skills_sentence = f"My expertise in {job_info['skills'][0]} aligns closely with the requirements for the {job_info['job_title']} role."
        else:
            skills_list = ', '.join(job_info['skills'][:-1])
            skills_sentence = f"My expertise in {skills_list}, and {job_info['skills'][-1]} aligns closely with the requirements for the {job_info['job_title']} role."
    else:
        skills_sentence = f"My expertise aligns closely with the requirements for the {job_info['job_title']} role."

    # Prepare the content by replacing placeholders
    cover_letter_content = template.format(
        your_name=applicant_info['name'],
        your_address=applicant_info['address'],
        your_city_state_zip=applicant_info['city_state_zip'],
        your_email=applicant_info['email'],
        your_phone=applicant_info['phone'],
        date=datetime.datetime.now().strftime('%B %d, %Y'),
        hiring_manager_name=job_info.get('hiring_manager_name', 'Hiring Manager'),
        company_name=job_info['company'],
        company_address=job_info.get('company_address', ''),
        company_city_state_zip=job_info.get('company_city_state_zip', ''),
        job_title=job_info['job_title'],
        skills_sentence=skills_sentence
    )

    # Save the cover letter as a Word document in the tailored_documents directory
    safe_company = job_info['company'].replace(' ', '_').replace('/', '_')
    safe_title = job_info['job_title'].replace(' ', '_').replace('/', '_')
    os.makedirs('tailored_documents', exist_ok=True)
    tailored_cover_letter_path = os.path.join('tailored_documents', f"Cover_Letter_{safe_company}_{safe_title}.docx")
    doc = Document()
    for line in cover_letter_content.strip().split('\n'):
        doc.add_paragraph(line)
    doc.save(tailored_cover_letter_path)
    return tailored_cover_letter_path


def prepare_application_data(df, resume_path, cover_letter_template, applicant_info, progress_dialog=None):
    applications = []
    total_jobs = len(df)
    current_progress = progress_dialog.value() if progress_dialog else 0
    for index, row in df.iterrows():
        # Clean the job summary
        cleaned_summary = clean_text(row['Summary'])
        # Extract keywords from the job description
        job_keywords = extract_keywords(cleaned_summary)
        job_info = prepare_job_info(row, job_keywords)
        # Generate tailored resume and cover letter
        tailored_resume = update_resume(resume_path, job_keywords, job_info)
        tailored_cover_letter = generate_cover_letter(
            cover_letter_template,
            applicant_info,
            job_info
        )
        # Store application data, including the extracted skills and default status
        applications.append({
            'Company': row['Company'],
            'Job Title': row['Title'],
            'Link': row['Link'],
            'Tailored Resume': tailored_resume,
            'Tailored Cover Letter': tailored_cover_letter,
            'Skills': ', '.join(job_keywords),
            'Status': 'Not Applied',  # Add default status
        })

        # Update progress
        if progress_dialog:
            current_progress += 1
            progress_dialog.setValue(current_progress)
            progress_dialog.setLabelText(f"Processing application {index + 1}/{total_jobs}")
            QApplication.processEvents()

    # Save all applications to Applications.csv
    applications_df = pd.DataFrame(applications)
    applications_df.to_csv('Applications.csv', index=False)
    print("Application data prepared and saved to Applications.csv.")


# Data Preparation Functions End

if __name__ == '__main__':
    app = QApplication(sys.argv)
    viewer = ApplicationViewer()
    viewer.show()
    sys.exit(app.exec_())
